import os
from io import BytesIO
from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from openai import OpenAI
from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from .forms import GeneratePostsForm
from content_agent.models import StrategyRequest

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def generate_posts(request):
    if request.method == 'POST':
        form = GeneratePostsForm(request.POST)
        if form.is_valid():
            strategy_id = form.cleaned_data['strategy_id']
            number_of_posts = form.cleaned_data['number_of_posts']
            name = form.cleaned_data['name']
            email = form.cleaned_data['email']

            strategy = get_object_or_404(StrategyRequest, id=strategy_id)

            print("\n===== GENERATE POSTS WITH MARKDOWN PROCESSING FOR .DOCX STARTED =====")

            import re
            titles = re.findall(r'^\s*\d+\.\s+(.*)', strategy.result, flags=re.MULTILINE)
            print("EXTRACTED TITLES:", titles)

            if not titles:
                fallback_titles = []
                for line in strategy.result.splitlines():
                    line = line.strip()
                    if (
                        line and
                        len(line) < 120 and
                        not line.lower().startswith(("cta", "month")) and
                        not line.endswith(":") and
                        line[0].isalpha()
                    ):
                        fallback_titles.append(line)
                titles = fallback_titles
                print("FALLBACK TITLES:", titles)

            if not titles:
                titles = ["Local SEO Tips for Cedar Rapids Businesses"]
                print("USING DEFAULT TITLE:", titles)

            titles = titles[:number_of_posts]
            print(f"TITLES TO GENERATE ({len(titles)}):", titles)

            doc = Document()
            doc.add_heading(f'Blog Posts for {name}', 0)
            doc.add_paragraph(f'Generated for: {email}\n')

            for idx, title in enumerate(titles, start=1):
                prompt = (
                    f"Write a detailed, SEO-friendly blog post draft (~800 words) for the following title:\n"
                    f"Title: {title}\n"
                    f"Ensure the tone matches: {strategy.tone}.\n"
                    f"The post should be helpful, well-structured with H2 and H3 headings."
                )

                try:
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "You are a professional content writer AI agent."},
                            {"role": "user", "content": prompt}
                        ]
                    )
                    content_markdown = response.choices[0].message.content

                    # Convert Markdown to HTML
                    html_content = markdown(content_markdown)
                    soup = BeautifulSoup(html_content, "html.parser")

                    # Insert into docx with basic formatting
                    doc.add_heading(f"{idx}. {title}", level=1)

                    for element in soup.contents:
                        if element.name == 'h1':
                            doc.add_heading(element.get_text(strip=True), level=1)
                        elif element.name == 'h2':
                            doc.add_heading(element.get_text(strip=True), level=2)
                        elif element.name == 'h3':
                            doc.add_heading(element.get_text(strip=True), level=3)
                        elif element.name == 'ul':
                            for li in element.find_all('li'):
                                p = doc.add_paragraph(style='ListBullet')
                                p.add_run(li.get_text(strip=True))
                        elif element.name == 'ol':
                            for li in element.find_all('li'):
                                p = doc.add_paragraph(style='ListNumber')
                                p.add_run(li.get_text(strip=True))
                        elif element.name == 'p':
                            doc.add_paragraph(element.get_text(strip=True))
                        elif element.name == 'blockquote':
                            p = doc.add_paragraph(element.get_text(strip=True))
                            p.style.font.italic = True
                        else:
                            if element.string and element.string.strip():
                                doc.add_paragraph(element.string.strip())

                    print(f"[{idx}] Post generated for title: {title}")

                except Exception as e:
                    error_msg = f"Error generating post for '{title}': {str(e)}"
                    doc.add_heading(f"{idx}. {title}", level=1)
                    doc.add_paragraph(error_msg)
                    print(error_msg)

            # Prepare file for download
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            safe_name = name.replace(" ", "_")
            response = HttpResponse(
                file_stream.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{safe_name}_blog_posts.docx"'
            return response

    else:
        form = GeneratePostsForm()

    return render(request, 'content_writer_agent/generate_posts.html', {
        'form': form
    })
